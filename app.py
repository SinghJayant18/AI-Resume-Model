import streamlit as st
import pdfplumber
from docx import Document
import io
import re
import os
from dotenv import load_dotenv
import google.generativeai as genai

# --------------------
# Gemini API Config
# --------------------
load_dotenv()  # Load environment variables from .env file
API_KEY = os.getenv("API_KEY")
genai.configure(api_key=API_KEY)
st.set_page_config(page_title="Smart Resume Tailor (Gemini)", layout="wide")
# Auto-detect latest Gemini model that supports generateContent
# Removed dynamic model fetching; using static model name
MODEL_NAME = "gemini-pro"
st.sidebar.info(f"Using model: {MODEL_NAME}")

# --------------------
# Streamlit UI Setup
# --------------------
st.title("Smart Resume Tailor — ATS-friendly Resume Optimizer (Gemini API)")
st.write("This app uses the Gemini API to help you tailor your resume to match job descriptions, optimizing for ATS systems.")  

uploaded_file = st.file_uploader("Upload resume (PDF, DOCX, TXT)", type=["pdf", "docx", "txt"])
job_desc = st.text_area("Paste job description here", height=200)

# --------------------
# Helper: Extract text
# --------------------
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    data = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    elif name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs if p.text])

    else:
        return data.decode("utf-8", errors="ignore")

# --------------------
# Helper: Gemini prompt
# --------------------
def gemini_prompt(system_prompt, user_prompt):
    model = genai.GenerativeModel(MODEL_NAME)
    prompt_text = f"{system_prompt}\n\n{user_prompt}"
    response = model.generate_content(prompt_text)
    return response.text.strip()

# --------------------
# Main logic
# --------------------
if st.button("Analyze & Tailor"):
    if not uploaded_file:
        st.error("Please upload a resume file.")
    elif not job_desc.strip():
        st.error("Please paste a job description.")
    else:
        with st.spinner("Extracting resume..."):
            resume_text = extract_text_from_file(uploaded_file)

        st.subheader("Extracted Resume Text")
        st.text_area("Resume", resume_text, height=200)

        # Single Gemini LLM call for ATS Score, Matched/Missing Skills, and Tailored Resume
        unified_prompt = f"""
Given the following ORIGINAL RESUME and JOB DESCRIPTION, perform the following:

1. Calculate the ATS match score (0-100%) based ONLY on the overlap of skillset and life experience.
2. Return ONLY:
   - The ATS match score as a percentage (e.g., 75%).
   - A comma-separated list of matched skills/experiences.
   - A comma-separated list of missing skills/experiences.
3. Then, as an expert resume writer, produce:
   - Dummy/fake projects that match the job description, including all tech stack and improvements made in the projects.
   - Suggestions for projects that can be added to the resume which match the job description, with proper details and dummy projects, tech stack, and improvements (e.g., reduced latency, increased throughput, etc.).
   - A short 2-3 line professional summary.
   - An ATS-optimized resume (sections: Name, Contact, Summary, Skills, Experience with bullet points, Education).
   - A bullet list of improvement suggestions.

ORIGINAL RESUME:
{resume_text}

JOB DESCRIPTION:
{job_desc}
"""
        unified_result = gemini_prompt("You are a helpful ATS and resume assistant.", unified_prompt)

        # Parse the unified_result for ATS score, matched, and missing (try to extract if possible)
        score = 0
        matched = []
        unmatched = []
        lines = unified_result.split("\n")
        for line in lines:
            if "match score" in line.lower():
                score = int(''.join(filter(str.isdigit, line)))
            elif "matched" in line.lower():
                matched = [k.strip() for k in line.split(":",1)[-1].split(",") if k.strip()]
            elif "missing" in line.lower():
                unmatched = [k.strip() for k in line.split(":",1)[-1].split(",") if k.strip()]
        st.metric("ATS Match Score", f"{score}%")
        st.write("Matched:", matched)
        st.write("Missing:", unmatched)

        st.subheader("Tailored Resume Output")
        st.write(unified_result)

        # Save DOCX for download
        doc_buffer = io.BytesIO()
        doc = Document()
        for line in unified_result.split("\n"):
            doc.add_paragraph(line)
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        st.download_button(
            label="Download Tailored Resume (DOCX)",
            data=doc_buffer,
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success("Resume tailored successfully!")
else:
    st.info("Upload a resume and paste a job description to get started.")

# --------------------
# Sidebar Info
# --------------------
st.sidebar.markdown("### About")
st.sidebar.markdown("This app helps you tailor your resume to match job descriptions using AI. It extracts keywords, calculates ATS scores, and generates a tailored resume.")
st.sidebar.markdown("### Contact")
st.sidebar.markdown("For any issues or feedback, please reach out to the developer.")
st.sidebar.markdown("### Disclaimer")
st.sidebar.markdown("This app is for educational purposes only. The results may vary based on the input data and the AI model's capabilities.")
st.sidebar.markdown("### Privacy")
st.sidebar.markdown("Your uploaded files and job descriptions are processed in memory and not stored. Please ensure you do not upload sensitive information.")
st.sidebar.markdown("### License")
st.sidebar.markdown("This app is open-source. You can find the source code on [GitHub](https://github.com/your-repo).")
